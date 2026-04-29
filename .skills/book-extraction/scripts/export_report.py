#!/usr/bin/env python3
"""
书籍深度榨取报告导出工具
支持格式：Markdown/PDF/Word

使用方法：
    python export_report.py <input_file> [options]

示例：
    python export_report.py book-report.md --format pdf
    python export_report.py book-report.md --format pdf,word
"""

import argparse
import re
from pathlib import Path
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from markdown import markdown
    from weasyprint import HTML
except ImportError as e:
    print(f"缺少依赖库: {e}")
    print("请安装依赖: pip install markdown python-docx weasyprint")
    sys.exit(1)


def export_to_pdf(input_file, output_file=None):
    """导出为PDF格式"""
    if output_file is None:
        output_file = Path(input_file).with_suffix('.pdf')
    else:
        output_file = Path(output_file)

    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 将Markdown转换为HTML
    html_content = markdown(md_content)

    # 添加基础CSS样式（使用更通用的字体栈，确保中文可显示）
    css = """
    <style>
    body {
        font-family: "Noto Sans CJK", "Source Han Sans", "Microsoft YaHei", "SimSun", "PingFang SC", "Heiti SC", sans-serif;
        line-height: 1.6;
        margin: 40px;
        font-size: 12pt;
    }
    h1 {
        color: #333;
        border-bottom: 2px solid #333;
        padding-bottom: 10px;
        font-size: 24pt;
    }
    h2 {
        color: #555;
        margin-top: 30px;
        font-size: 18pt;
    }
    h3 {
        color: #666;
        font-size: 14pt;
    }
    table {
        border-collapse: collapse;
        width: 100%;
        margin: 20px 0;
        font-size: 10pt;
    }
    th, td {
        border: 1px solid #ddd;
        padding: 8px;
        text-align: left;
    }
    th {
        background-color: #f2f2f2;
        font-weight: bold;
    }
    hr {
        border: none;
        border-top: 1px solid #eee;
        margin: 20px 0;
    }
    ul, ol {
        margin: 10px 0;
        padding-left: 20px;
    }
    li {
        margin: 5px 0;
    }
    strong {
        font-weight: bold;
    }
    em {
        font-style: italic;
    }
    code {
        font-family: monospace;
        background-color: #f5f5f5;
        padding: 2px 6px;
        border-radius: 3px;
        font-size: 10pt;
    }
    </style>
    """

    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        {css}
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    try:
        # 使用weasyprint生成PDF
        HTML(string=full_html).write_pdf(output_file)
        print(f"✓ PDF格式已导出: {output_file}")
        return str(output_file)
    except Exception as e:
        print(f"✗ PDF导出失败: {e}")
        print("提示: 如果问题持续，建议使用Word格式导出后再转换为PDF")
        raise


def export_to_word(input_file, output_file=None):
    """导出为Word格式"""
    if output_file is None:
        output_file = Path(input_file).with_suffix('.docx')
    else:
        output_file = Path(output_file)

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # 标题1
        if line.startswith('# '):
            heading = line[2:].strip()
            doc.add_heading(heading, level=1)
        # 标题2
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
        # 标题3
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=3)
        # 标题4
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=4)
        # 列表项
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
        # 表格（简化处理，只添加文本）
        elif line.startswith('|'):
            p = doc.add_paragraph(line.strip())
            p.runs[0].font.size = Pt(9)
        # 分割线
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        # 普通段落
        else:
            # 处理加粗
            if '**' in line:
                # 简化处理：移除markdown标记
                clean_line = line.replace('**', '')
                doc.add_paragraph(clean_line)
            else:
                doc.add_paragraph(line)

    doc.save(output_file)

    print(f"✓ Word格式已导出: {output_file}")
    return str(output_file)


def export_to_formats(input_file, formats, output_dir=None, base_name=None):
    """
    批量导出多种格式

    参数:
        input_file: 输入的Markdown文件路径
        formats: 要导出的格式列表，如 ['pdf', 'word']
        output_dir: 输出目录（可选）
        base_name: 输出文件基础名（可选，不包含扩展名）

    返回:
        成功导出的文件列表
    """
    if base_name is None:
        base_name = Path(input_file).stem

    if output_dir is None:
        output_dir = Path(input_file).parent
    else:
        output_dir = Path(output_dir)

    export_functions = {
        'pdf': export_to_pdf,
        'word': export_to_word
    }

    results = {
        'success': [],
        'failed': []
    }

    for fmt in formats:
        if fmt not in export_functions:
            print(f"✗ 不支持的格式: {fmt}")
            results['failed'].append({'format': fmt, 'error': '不支持的格式'})
            continue

        try:
            # 处理文件扩展名
            if fmt == 'word':
                ext = 'docx'
            else:
                ext = fmt
            output_file = output_dir / f"{base_name}.{ext}"
            export_functions[fmt](input_file, str(output_file))
            results['success'].append(str(output_file))
        except Exception as e:
            print(f"✗ 导出{fmt}格式失败: {e}")
            results['failed'].append({'format': fmt, 'error': str(e)})

    return results


def main():
    parser = argparse.ArgumentParser(
        description='书籍深度榨取报告导出工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 预览Markdown报告（不导出）
  python export_report.py book-report.md

  # 导出为PDF格式
  python export_report.py book-report.md --format pdf

  # 导出为Word格式
  python export_report.py book-report.md --format word

  # 同时导出PDF和Word（用户可自主选择需要的格式）
  python export_report.py book-report.md --format pdf,word

  # 导出到指定目录
  python export_report.py book-report.md --format pdf,word --output-dir ./output
        """
    )
    parser.add_argument('input_file', help='输入的Markdown报告文件')
    parser.add_argument('--format', help='输出格式，可选值：markdown/pdf/word，多个格式用逗号分隔（如：pdf,word）')
    parser.add_argument('--output', '-o', help='输出文件路径（可选，仅在单格式导出时有效）')
    parser.add_argument('--output-dir', help='输出目录（可选，批量导出时使用）')
    parser.add_argument('--base-name', help='输出文件基础名（可选，不包含扩展名）')

    args = parser.parse_args()

    # 检查输入文件是否存在
    if not Path(args.input_file).exists():
        print(f"错误: 输入文件不存在: {args.input_file}")
        sys.exit(1)

    # 如果没有指定格式，只显示Markdown预览
    if not args.format:
        print("=== 📖 Markdown报告预览 ===\n")
        with open(args.input_file, 'r', encoding='utf-8') as f:
            content = f.read()
            # 只显示前1500字符作为预览
            preview = content if len(content) <= 1500 else content[:1500] + "\n\n... (报告太长，已截断预览) ..."
            print(preview)
        print(f"\n{'='*60}")
        print(f"完整报告已保存到: {Path(args.input_file).absolute()}")
        print(f"{'='*60}")
        print("\n💡 如需导出其他格式，请使用以下命令：")
        print(f"  python export_report.py {args.input_file} --format pdf")
        print(f"  python export_report.py {args.input_file} --format word")
        print(f"  python export_report.py {args.input_file} --format pdf,word")
        return

    # 处理导出格式
    formats = [f.strip().lower() for f in args.format.split(',')]

    # 去重
    formats = list(set(formats))

    if len(formats) > 1:
        # 批量导出
        print(f"\n{'='*60}")
        print(f"=== 🚀 开始批量导出 {len(formats)} 种格式 ===")
        print(f"{'='*60}\n")
        results = export_to_formats(args.input_file, formats, args.output_dir, args.base_name)

        print(f"\n{'='*60}")
        print(f"=== ✅ 导出完成 ===")
        print(f"{'='*60}")
        print(f"\n✓ 成功: {len(results['success'])} 个")
        for f in results['success']:
            print(f"  - {f}")

        if results['failed']:
            print(f"\n✗ 失败: {len(results['failed'])} 个")
            for item in results['failed']:
                print(f"  - {item['format']}: {item['error']}")
    else:
        # 单格式导出
        export_functions = {
            'markdown': lambda f, o: print(f"✓ Markdown已存在于: {f}") if not o else None,
            'pdf': export_to_pdf,
            'word': export_to_word
        }
        export_func = export_functions.get(formats[0] if formats else 'markdown')
        if export_func:
            try:
                if formats[0] == 'markdown':
                    export_func(args.input_file, args.output)
                else:
                    export_func(args.input_file, args.output)
                print(f"\n✓ 导出完成！")
            except Exception as e:
                print(f"错误: 导出失败 - {e}")
                import traceback
                traceback.print_exc()
                sys.exit(1)


if __name__ == '__main__':
    main()
